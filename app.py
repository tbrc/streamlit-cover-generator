import os
import re
import io
import zipfile
import base64
import streamlit as st
import pandas as pd
from datetime import datetime
import datetime
from PIL import Image
from dotenv import load_dotenv

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml, ns
from docx.oxml.ns import nsdecls

# --------------------------------------------------
# INITIAL SETUP
# --------------------------------------------------
if os.path.exists(".env"):
    load_dotenv()

client = OpenAI(api_key=os.getenv("COVER_IMAGES_API_KEY"))

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_IMG = os.path.join(BASE_DIR, "temp/images")
TEMP_DOC = os.path.join(BASE_DIR, "temp/docs")      
ASSETS = os.path.join(BASE_DIR, "assets")

DESIGN_PATH = os.path.join(ASSETS, "design.png")
LOGO_PATH = os.path.join(ASSETS, "logo.png")

os.makedirs(TEMP_IMG, exist_ok=True)
os.makedirs(TEMP_DOC, exist_ok=True)

# --------------------------------------------------
# HELPERS
# --------------------------------------------------
def clean_market_name(name):
    return re.sub(r"\s+Report\s+\d{4}$", "", name).strip()

def make_image_bottom_aligned(run):
    drawing = run._r.xpath('.//w:drawing')[0]
    inline = drawing.xpath('.//wp:inline')[0]

    anchor_xml = f'''
    <wp:anchor {nsdecls("wp", "a", "r")}
        simplePos="0"
        relativeHeight="0"
        behindDoc="1"
        locked="0"
        layoutInCell="1"
        allowOverlap="1">

        <wp:simplePos x="0" y="0"/>

        <wp:positionH relativeFrom="page">
            <wp:align>center</wp:align>
        </wp:positionH>

        <wp:positionV relativeFrom="page">
            <wp:align>bottom</wp:align>
        </wp:positionV>

        <wp:extent cx="{inline.extent.cx}" cy="{inline.extent.cy}"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="BottomImage"/>
        <wp:cNvGraphicFramePr/>
        {inline.graphic.xml}
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    drawing.replace(inline, anchor)     

def set_table_width(table, width_in_inches):
    tblW = OxmlElement('w:tblW')
    tblW.set(ns.qn('w:w'), str(int(width_in_inches * 1440)))
    tblW.set(ns.qn('w:type'), 'dxa')
    table._tbl.tblPr.append(tblW)

# --------------------------------------------------
# IMAGE GENERATION
# --------------------------------------------------
def generate_cover_image(market):
    prompt = (
        f"Generate a cover image in 16:9 ratio, 1200 pixels wide, on the {market}. No text needed. Only need real images, avoid illustrational images. Do not combine multiple visual elements into a single image. No collages, no layered imagery, no surreal or conceptual blending—keep the image simple, realistic, and immediately understandable. Ensure the image is a single unified scene. Avoid divided layouts, side-by-side compositions, split screens, panels, or any visually distinct sections."
    )

    response = client.images.generate(
        model="gpt-image-1",
        prompt=prompt,
        size="1536x1024"
    )

    img_bytes = base64.b64decode(response.data[0].b64_json)
    market_img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
    design = Image.open(DESIGN_PATH).convert("RGBA")

    scale = 0.825
    market_img = market_img.resize(
        (int(design.width * scale), int(design.height * scale)),
        Image.LANCZOS
    )

    canvas = Image.new("RGBA", design.size, (0, 0, 0, 0))
    x = (design.width - market_img.width) // 2 + 136
    y = (design.height - market_img.height) // 2 + 13

    canvas.paste(market_img, (x, y), market_img)
    final = Image.alpha_composite(canvas, design)

    out_path = os.path.join(TEMP_IMG, f"{market}.png")
    final.save(out_path)

    return out_path

# --------------------------------------------------
# DOCX CREATION
# --------------------------------------------------
def create_cover_docx(market, date_str, code, image_path):
    doc = Document()

    # -----------------------------
    # PAGE MARGINS = 0
    # -----------------------------
    section = doc.sections[0]
    section.top_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.footer_distance = Inches(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Inches(1)

    table = doc.add_table(rows=2, cols=2)
    table.allow_autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Reduce table width
    set_table_width(table, 7.0)
    table.columns[0].width = Inches(4.2)
    table.columns[1].width = Inches(2.8)

    # Row 1: Market Name + Logo
    cell1 = table.cell(0, 0)
    para1 = cell1.paragraphs[0]
    run1 = para1.add_run(market)
    run1.font.name = 'Roboto'
    run1.bold = True
    run1.font.color.rgb = RGBColor(36, 57, 78)  # #24394e
    
    # Adjust font size based on character length
    market_len = len(market)
    if market_len <= 65:
        run1.font.size = Pt(30)
    elif market_len <= 105:
        run1.font.size = Pt(24)
    else:
        run1.font.size = Pt(20)
    
    para1.alignment = WD_ALIGN_PARAGRAPH.LEFT


    cell2 = table.cell(0, 1)
    para2 = cell2.paragraphs[0]
    run2 = para2.add_run()
    run2.add_picture(LOGO_PATH, width=Inches(2.00), height=Inches(0.64))
    para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 2: Date + Code
    date_cell = table.cell(1, 0)
    date_para = date_cell.paragraphs[0]
    date_para.paragraph_format.space_before = Pt(20)
    run_date_label = date_para.add_run("PUBLISHED DATE:")
    run_date_label.font.name = 'Roboto'
    run_date_label.font.size = Pt(15)
    run_date_label.bold = True
    run_date_label.font.color.rgb = RGBColor(36, 57, 78)  # #24394e
    date_para.add_run("\n")
    run_date_value = date_para.add_run(date_str)
    run_date_value.font.name = 'Roboto'
    run_date_value.font.size = Pt(17)
    run_date_value.font.color.rgb = RGBColor(33, 123, 108)  # #217b6c
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    code_cell = table.cell(1, 1)
    code_para = code_cell.paragraphs[0]
    code_para.paragraph_format.space_before = Pt(20)
    run_code_label = code_para.add_run("CODE:")
    run_code_label.font.name = 'Roboto'
    run_code_label.font.size = Pt(15)
    run_code_label.bold = True
    run_code_label.font.color.rgb = RGBColor(36, 57, 78)  # #24394e
    code_para.add_run("\n")
    run_code_value = code_para.add_run(code)
    run_code_value.font.name = 'Roboto'
    run_code_value.font.size = Pt(17)
    run_code_value.font.color.rgb = RGBColor(33, 123, 108)  # #217b6c
    code_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -----------------------------
    # IMAGE flush at end of page
    # -----------------------------
    img_para = doc.add_paragraph()
    img_para.paragraph_format.space_before = Pt(0)
    img_para.paragraph_format.space_after = Pt(0)
    
    page_width = section.page_width.inches
    run_img = img_para.add_run()
    run_img.add_picture(image_path, width=Inches(page_width))
    
    # 🔥 THIS LINE DOES THE MAGIC
    make_image_bottom_aligned(run_img)

    out_path = os.path.join(TEMP_DOC, f"{market}.docx")
    doc.save(out_path)

    return out_path

# --------------------------------------------------
# STREAMLIT UI
# --------------------------------------------------
st.set_page_config("Cover Page Generator", layout="centered")
st.title("📘 Market Report Cover Page Generator")

excel_file = st.file_uploader("Upload Excel File", type=["xlsx"])

if excel_file and st.button("Generate Covers"):
    df = pd.read_excel(excel_file)

    progress = st.progress(0)
    docs = []

    status_box = st.empty()
    status_messages = []

    for i, row in df.iterrows():
        original_market = str(row["Product Name"]).strip()   # EXACT Excel value
        image_market = clean_market_name(original_market)    # Only for image

        date = row["Published Date"]
        code = str(row["Report Code"])

        dt = datetime.datetime.strptime(date, "%A, %B %d, %Y")

        date_str = dt.strftime("%d %B %Y")

        img = generate_cover_image(image_market)
        docx = create_cover_docx(original_market, date_str, code, img)
        docs.append(docx)

        # ✅ STATUS UPDATE AFTER COMPLETION
        status_messages.append(f"✅ DOCX generated for **{original_market}**")
        status_box.markdown("<br>".join(status_messages), unsafe_allow_html=True)

        progress.progress((i + 1) / len(df))

    zip_path = os.path.join(BASE_DIR, "cover_pages.zip")
    with zipfile.ZipFile(zip_path, "w") as z:
        for d in docs:
            z.write(d, arcname=os.path.basename(d))

    with open(zip_path, "rb") as f:
        st.download_button("⬇️ Download ZIP", f, file_name="cover_pages.zip")





